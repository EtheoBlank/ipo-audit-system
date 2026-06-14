"""Document parser for the sales-ledger module.

Accepts UploadFile (docx / pdf / xlsx) and returns a (doc_type, raw_text) pair
where raw_text is a Markdown-ish representation suitable for prompting an LLM.

Parsing strategy:
  - xlsx/xls: every sheet is converted to a Markdown table via pandas.
  - docx:     python-docx walks paragraphs and tables, preserving structure.
  - pdf:      pdfplumber extracts text + tables page by page.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Tuple

import pandas as pd
from fastapi import UploadFile

logger = logging.getLogger(__name__)

SUPPORTED_EXTS = {".docx", ".pdf", ".xlsx", ".xls"}


class DocumentParserError(RuntimeError):
    """Raised when a document cannot be parsed."""


class DocumentParser:
    """Static helpers — no shared state needed."""

    @staticmethod
    def ext_of(filename: str) -> str:
        return Path(filename).suffix.lower()

    @classmethod
    async def parse(cls, upload: UploadFile, save_dir: Path) -> Tuple[str, str]:
        """Parse an uploaded file and return (doc_type, raw_text)."""
        ext = cls.ext_of(upload.filename or "")
        if ext not in SUPPORTED_EXTS:
            raise DocumentParserError(f"不支持的文件类型: {ext}。仅支持 {sorted(SUPPORTED_EXTS)}")

        save_dir.mkdir(parents=True, exist_ok=True)
        # Persist to disk so the various parsers can re-open the file.
        temp_path = save_dir / f"sales_src_{upload.filename}"
        content = await upload.read()
        temp_path.write_bytes(content)

        try:
            if ext in {".xlsx", ".xls"}:
                return ext.lstrip("."), cls._parse_xlsx(temp_path)
            if ext == ".docx":
                return ext.lstrip("."), cls._parse_docx(temp_path)
            if ext == ".pdf":
                return ext.lstrip("."), cls._parse_pdf(temp_path)
        except DocumentParserError:
            raise
        except Exception as exc:  # noqa: BLE001 — surface a friendly error
            logger.exception("Parse failure for %s", upload.filename)
            raise DocumentParserError(f"解析 {upload.filename} 失败: {exc}") from exc
        finally:
            try:
                temp_path.unlink(missing_ok=True)
            except Exception:  # noqa: BLE001
                pass

        raise DocumentParserError(f"未知文件类型: {ext}")

    # --- format-specific helpers -----------------------------------------

    @staticmethod
    def _parse_xlsx(path: Path) -> str:
        """Read every sheet, render each as a Markdown table."""
        try:
            sheets = pd.read_excel(path, sheet_name=None, dtype=str)
        except Exception as exc:  # noqa: BLE001
            raise DocumentParserError(f"无法读取 Excel: {exc}") from exc

        chunks: list[str] = []
        for name, df in sheets.items():
            chunks.append(f"## Sheet: {name}\n")
            if df.empty:
                chunks.append("(空)\n")
                continue
            df = df.fillna("")
            chunks.append(df.to_markdown(index=False))
            chunks.append("\n")
        return "\n".join(chunks)

    @staticmethod
    def _parse_docx(path: Path) -> str:
        try:
            from docx import Document  # python-docx
        except ImportError as exc:  # pragma: no cover
            raise DocumentParserError("缺少 python-docx 依赖") from exc

        doc = Document(str(path))
        chunks: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                chunks.append(text)
        for ti, table in enumerate(doc.tables, 1):
            chunks.append(f"\n## Table {ti}\n")
            rows: list[list[str]] = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if not rows:
                continue
            header, *body = rows
            df = pd.DataFrame(body, columns=header)
            chunks.append(df.to_markdown(index=False))
            chunks.append("")
        return "\n".join(chunks)

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        try:
            import pdfplumber
        except ImportError as exc:  # pragma: no cover
            raise DocumentParserError("缺少 pdfplumber 依赖，请先 `uv add pdfplumber`") from exc

        chunks: list[str] = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                chunks.append(f"## Page {i}\n")
                text = page.extract_text() or ""
                if text.strip():
                    chunks.append(text)
                for ti, table in enumerate(page.extract_tables() or [], 1):
                    if not table:
                        continue
                    header, *body = table
                    # pdfplumber returns None inside cells; coerce to "".
                    header = [("" if c is None else str(c)) for c in header]
                    body = [[("" if c is None else str(c)) for c in r] for r in body]
                    df = pd.DataFrame(body, columns=header)
                    chunks.append(f"\n### Page {i} Table {ti}\n")
                    chunks.append(df.to_markdown(index=False))
                chunks.append("")
        return "\n".join(chunks)
