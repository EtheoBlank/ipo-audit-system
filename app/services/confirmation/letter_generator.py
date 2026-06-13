"""Confirmation letter generator (询证函生成器).

生成银行 / 客户 / 供应商 / 其他往来询证函，遵循:
- 银行询证函: 财政部《银行询证函参考格式》(财会[2024]6号 等更新)
- 其他往来函证: 《中国注册会计师审计准则第 1311 号 — 对存货、诉讼和索赔的
  审计》、《第 1504 号 — 在审计报告中沟通关键审计事项》及问题解答
  —— 要求至少函证: 余额 + 本期发生额 + 关键合同条款 + 已背书票据

输出格式: docx (默认) / pdf (经由 docx2pdf，可选)
发函内容以 JSON 快照固化在 ConfirmationLetter.content_snapshot，
发函日期/金额在 ConfirmationLetter.amount_snapshot 固化。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)


class LetterGenerationError(RuntimeError):
    """Raised when a confirmation letter cannot be generated."""


# ---- 银行询证函官方模板 (财会[2024]6号 风格) -----------------------


BANK_LETTER_TEMPLATE = """\
{cpa_firm}
──────────────────────────────────────
询 证 函
（银行询证函参考格式）

{recipient}：

本公司聘请的 {cpa_firm} 正在对本公司 {period} 财务报表进行审计，按照中国
注册会计师审计准则的要求，应当询证本公司与贵行的相关业务信息。下列信息
出自本公司账簿记录，如与贵行记录相符，请在本函下端『信息证明无误』处签
章证明；如有不符，请在『信息不符』处列示不符金额及相关说明。本函仅为复
核账目之用，并非催款结算。

回函请直接寄至：
{cpa_firm} {cpa_address}
收件人：{auditor_name}  电话：{auditor_phone}  邮编：{auditor_zip}

──────────────────────────────────────
一、本公司截至 {balance_date} 在贵行的存款、贷款及有关业务情况

【存款】
  1. 活期存款余额    RMB {current_deposit:,.2f}    存款利率 {current_deposit_rate}
  2. 定期存款余额    RMB {time_deposit:,.2f}      存款利率 {time_deposit_rate}
  3. 通知存款余额    RMB {notice_deposit:,.2f}
  4. 协定存款余额    RMB {agreement_deposit:,.2f}
  5. 结构性存款      RMB {structured_deposit:,.2f}
  6. 外币存款折 RMB   {fx_deposit:,.2f}（原币 {fx_deposit_currency}）
  □ 上述存款是否存在质押/冻结/担保：{deposit_pledge_status}

【贷款】
  7. 短期借款余额    RMB {short_loan:,.2f}        利率 {short_loan_rate}
  8. 长期借款余额    RMB {long_loan:,.2f}         利率 {long_loan_rate}
  9. 应付债券余额    RMB {bond_payable:,.2f}      利率 {bond_rate}
 10. 委托贷款本金    RMB {entrusted_loan:,.2f}    利率 {entrusted_loan_rate}
 11. 未偿还利息      RMB {interest_payable:,.2f}
 12. 授信额度        RMB {credit_line:,.2f}        已占用 {credit_used:,.2f}
 □ 上述贷款担保/抵押/质押情况：{loan_collateral_status}
 □ 是否存在违约/关注类/逾期：{loan_default_status}

【银行承兑汇票】
 13. 已开立银行承兑汇票余额  RMB {bankers_acceptance:,.2f}
 14. 商业承兑汇票余额        RMB {commercial_paper:,.2f}
 15. 已贴现未到期银行承兑汇票  RMB {discounted_ba:,.2f}
 16. 已背书未到期银行承兑汇票  RMB {endorsed_ba:,.2f}
 17. 票据质押情况            {note_pledge_status}
 18. 是否存在逾期/拒付票据    {note_overdue_status}

【信用证】
 19. 信用证余额（开立/未使用）RMB {lc_balance:,.2f}   到期日 {lc_expiry}

【保函】
 20. 履约保函余额    RMB {perf_guarantee:,.2f}
 21. 投标保函余额    RMB {bid_guarantee:,.2f}
 22. 预付款保函余额  RMB {prepay_guarantee:,.2f}
 23. 质保保函余额    RMB {warranty_guarantee:,.2f}

【对外担保】
 24. 本公司对外提供担保余额  RMB {external_guarantee:,.2f}
     被担保方：{guarantee_recipient}
     反担保安排：{counter_guarantee}

【其他】
 25. 资金归集/资金池账户余额  RMB {cash_pool:,.2f}
 26. 其他业务说明：{other_business}

──────────────────────────────────────
二、回函

【信息证明无误】
贵行确认上述 1-26 项内容均与本行记录相符。
                                                 银行业务章：
                                                 经办人：
                                                 日期：    年    月    日

【信息不符】
请列示不符项目与金额（可另附详单）：
_________________________________________________________________

                                                 银行业务章：
                                                 经办人：
                                                 日期：    年    月    日

──────────────────────────────────────
（本函仅为复核账目之用，请加盖银行业务章后寄回）

{company_name}（公章）
日期：{sent_date}
"""


# ---- 客户 / 供应商 询证函 -----------------------------------------


CUSTOMER_SUPPLIER_LETTER_TEMPLATE = """\
{cpa_firm}
──────────────────────────────────────
企 业 询 证 函
（适用于 {direction}）

{customer_company}：

本公司聘请的 {cpa_firm} 正在对本公司 {period} 财务报表进行审计，按照《中国
注册会计师审计准则第 1311 号 — 对存货、诉讼和索赔的审计》《第 1502 号 —
在审计报告中发表非无保留意见》《第 1504 号 — 在审计报告中沟通关键审计
事项》等准则的要求，应当询证本公司与贵公司的往来账项及相关业务信息。下列
数据出自本公司账簿记录，请贵公司核对后在本函下端『信息证明无误』处签章证
明，或在『信息不符』处列示不符金额及原因。如有贵公司已记录但本公司未记录
的业务（如在途商品、未结算费用等），亦请一并说明。

回函请直接寄至：
{cpa_firm} {cpa_address}
收件人：{auditor_name}  电话：{auditor_phone}

──────────────────────────────────────
一、余额信息（截至 {balance_date}）

  1. {direction}余额（按本公司账面）  RMB {book_balance:,.2f}
  2. 未结算发票明细  共 {unsettled_invoice_count} 笔，金额合计 RMB {unsettled_invoice_amount:,.2f}
       详见附表 1《未结算发票明细》

──────────────────────────────────────
二、本期发生额（{period_start} 至 {balance_date}）

  3. 本期{transaction_verb}额  RMB {transaction_amount:,.2f}
  4. 本期{repayment_verb}额  RMB {repayment_amount:,.2f}
       详见附表 2《本期交易明细》

──────────────────────────────────────
三、票据及或有事项

  5. 贵公司已背书给本公司的未到期票据  RMB {endorsed_to_us:,.2f}  共 {endorsed_to_us_count} 笔
  6. 本公司已背书给贵公司的未到期票据  RMB {endorsed_to_them:,.2f}  共 {endorsed_to_them_count} 笔
  7. 双方互为担保的余额  RMB {mutual_guarantee:,.2f}
  8. 是否存在未决诉讼/质量索赔/退货折让安排：{litigation_status}

──────────────────────────────────────
四、关键合同条款（请贵公司确认或更正）

  9. 主合同编号：{contract_no}
 10. 合同期间：{contract_period}
 11. {direction}条件：{payment_terms}
 12. 信用额度：RMB {credit_limit:,.2f}
 13. 所有权保留/担保安排：{title_retention}
 14. 争议解决方式：{dispute_resolution}
 15. 重要变更与补充协议：{contract_amendments}

──────────────────────────────────────
五、回函

【信息证明无误】
贵公司确认上述 1-15 项内容与本方记录相符。
                                       贵公司财务专用章：
                                       经办人：
                                       日期：    年    月    日

【信息不符】
请列示不符项目与金额（可另附详单）：
_________________________________________________________________

                                       贵公司财务专用章：
                                       经办人：
                                       日期：    年    月    日

──────────────────────────────────────
（本函仅为复核账目之用，请加盖公章后寄回）

{company_name}（公章）
日期：{sent_date}
"""


# ---- 其他往来 询证函（其他应收 / 其他应付） ----------------------


OTHER_RECEIVABLE_PAYABLE_TEMPLATE = """\
{cpa_firm}
──────────────────────────────────────
其 他 往 来 询 证 函

{recipient}：

本公司聘请的 {cpa_firm} 正在对本公司 {period} 财务报表进行审计，按照
中国注册会计师审计准则的要求，应当询证本公司与贵方/贵公司的其他往来
款项。下列数据出自本公司账簿记录，请核对后在本函下端签章证明。

──────────────────────────────────────
一、款项信息（截至 {balance_date}）

  1. 款项性质：{nature}
  2. 账面余额：RMB {book_balance:,.2f}
  3. 款项起讫日：{start_date} 至 {end_date}
  4. 本期发生额：{period_amount:,.2f}

二、其他信息
  5. 是否涉及关联方：{related_party}
  6. 预计可收回/需支付：{recoverability}
  7. 重要补充说明：{note}

──────────────────────────────────────
三、回函

【信息证明无误】
                                       {recipient}（公章）：
                                       经办人：
                                       日期：    年    月    日

【信息不符】
_________________________________________________________________

                                       {recipient}（公章）：
                                       经办人：
                                       日期：    年    月    日

{company_name}（公章）
日期：{sent_date}
"""


# ============================================================
#  生成器
# ============================================================


class ConfirmationLetterGenerator:
    """询证函生成器 — 渲染文本 / docx / pdf。"""

    def __init__(self, output_dir: Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ---- 渲染文本 -----------------------------------------------------

    def render_text(
        self,
        template_id: str,
        *,
        company_name: str,
        period: str,
        period_start: str,
        balance_date: str,
        sent_date: str,
        recipient: str,
        cpa_firm: str,
        cpa_address: str,
        auditor_name: str,
        auditor_phone: str,
        auditor_zip: str = "",
        book_balance: float = 0.0,
        # bank-specific (all defaulted)
        current_deposit: float = 0.0,
        time_deposit: float = 0.0,
        notice_deposit: float = 0.0,
        agreement_deposit: float = 0.0,
        structured_deposit: float = 0.0,
        fx_deposit: float = 0.0,
        fx_deposit_currency: str = "",
        deposit_pledge_status: str = "无",
        short_loan: float = 0.0,
        long_loan: float = 0.0,
        bond_payable: float = 0.0,
        entrusted_loan: float = 0.0,
        interest_payable: float = 0.0,
        credit_line: float = 0.0,
        credit_used: float = 0.0,
        loan_collateral_status: str = "无",
        loan_default_status: str = "无",
        bankers_acceptance: float = 0.0,
        commercial_paper: float = 0.0,
        discounted_ba: float = 0.0,
        endorsed_ba: float = 0.0,
        note_pledge_status: str = "无",
        note_overdue_status: str = "无",
        lc_balance: float = 0.0,
        lc_expiry: str = "",
        perf_guarantee: float = 0.0,
        bid_guarantee: float = 0.0,
        prepay_guarantee: float = 0.0,
        warranty_guarantee: float = 0.0,
        external_guarantee: float = 0.0,
        guarantee_recipient: str = "",
        counter_guarantee: str = "无",
        cash_pool: float = 0.0,
        other_business: str = "无",
        current_deposit_rate: str = "0.35%",
        time_deposit_rate: str = "1.50%",
        short_loan_rate: str = "3.45%",
        long_loan_rate: str = "3.85%",
        bond_rate: str = "4.20%",
        entrusted_loan_rate: str = "0.00%",
        # customer / supplier
        direction: str = "应收账款",  # 应收账款 / 应付账款
        transaction_verb: str = "销售",  # 销售 / 采购
        repayment_verb: str = "回款",  # 回款 / 付款
        transaction_amount: float = 0.0,
        repayment_amount: float = 0.0,
        unsettled_invoice_count: int = 0,
        unsettled_invoice_amount: float = 0.0,
        endorsed_to_us: float = 0.0,
        endorsed_to_us_count: int = 0,
        endorsed_to_them: float = 0.0,
        endorsed_to_them_count: int = 0,
        mutual_guarantee: float = 0.0,
        litigation_status: str = "无",
        contract_no: str = "",
        contract_period: str = "",
        payment_terms: str = "",
        credit_limit: float = 0.0,
        title_retention: str = "无",
        dispute_resolution: str = "诉讼至本公司所在地法院",
        contract_amendments: str = "无",
        # other receivable/payable
        nature: str = "",
        start_date: str = "",
        end_date: str = "",
        period_amount: float = 0.0,
        related_party: str = "否",
        recoverability: str = "",
        note: str = "",
    ) -> str:
        ctx = {
            "cpa_firm": cpa_firm,
            "recipient": recipient,
            "company_name": company_name,
            "period": period,
            "period_start": period_start,
            "balance_date": balance_date,
            "sent_date": sent_date,
            "cpa_address": cpa_address,
            "auditor_name": auditor_name,
            "auditor_phone": auditor_phone,
            "auditor_zip": auditor_zip,
            "book_balance": book_balance,
            "current_deposit": current_deposit,
            "time_deposit": time_deposit,
            "notice_deposit": notice_deposit,
            "agreement_deposit": agreement_deposit,
            "structured_deposit": structured_deposit,
            "fx_deposit": fx_deposit,
            "fx_deposit_currency": fx_deposit_currency,
            "deposit_pledge_status": deposit_pledge_status,
            "short_loan": short_loan,
            "long_loan": long_loan,
            "bond_payable": bond_payable,
            "entrusted_loan": entrusted_loan,
            "interest_payable": interest_payable,
            "credit_line": credit_line,
            "credit_used": credit_used,
            "loan_collateral_status": loan_collateral_status,
            "loan_default_status": loan_default_status,
            "bankers_acceptance": bankers_acceptance,
            "commercial_paper": commercial_paper,
            "discounted_ba": discounted_ba,
            "endorsed_ba": endorsed_ba,
            "note_pledge_status": note_pledge_status,
            "note_overdue_status": note_overdue_status,
            "lc_balance": lc_balance,
            "lc_expiry": lc_expiry,
            "perf_guarantee": perf_guarantee,
            "bid_guarantee": bid_guarantee,
            "prepay_guarantee": prepay_guarantee,
            "warranty_guarantee": warranty_guarantee,
            "external_guarantee": external_guarantee,
            "guarantee_recipient": guarantee_recipient,
            "counter_guarantee": counter_guarantee,
            "cash_pool": cash_pool,
            "other_business": other_business,
            "current_deposit_rate": current_deposit_rate,
            "time_deposit_rate": time_deposit_rate,
            "short_loan_rate": short_loan_rate,
            "long_loan_rate": long_loan_rate,
            "bond_rate": bond_rate,
            "entrusted_loan_rate": entrusted_loan_rate,
            "direction": direction,
            "transaction_verb": transaction_verb,
            "repayment_verb": repayment_verb,
            "transaction_amount": transaction_amount,
            "repayment_amount": repayment_amount,
            "unsettled_invoice_count": unsettled_invoice_count,
            "unsettled_invoice_amount": unsettled_invoice_amount,
            "endorsed_to_us": endorsed_to_us,
            "endorsed_to_us_count": endorsed_to_us_count,
            "endorsed_to_them": endorsed_to_them,
            "endorsed_to_them_count": endorsed_to_them_count,
            "mutual_guarantee": mutual_guarantee,
            "litigation_status": litigation_status,
            "contract_no": contract_no,
            "contract_period": contract_period,
            "payment_terms": payment_terms,
            "credit_limit": credit_limit,
            "title_retention": title_retention,
            "dispute_resolution": dispute_resolution,
            "contract_amendments": contract_amendments,
            "nature": nature,
            "start_date": start_date,
            "end_date": end_date,
            "period_amount": period_amount,
            "related_party": related_party,
            "recoverability": recoverability,
            "note": note,
        }

        if template_id == "bank_official":
            return BANK_LETTER_TEMPLATE.format(**ctx)
        if template_id in ("customer_std", "supplier_std"):
            # 根据 amount 符号自动调整 direction
            if book_balance < 0 and direction == "应收账款":
                direction = "预收账款"
            if book_balance > 0 and direction == "应付账款":
                pass
            ctx["direction"] = direction
            # P0 修复: 补齐 customer_company 等模板占位符, 避免 KeyError
            ctx.setdefault("customer_company", recipient)
            ctx.setdefault("endorsed_to_us", "0.00")
            ctx.setdefault("endorsed_to_us_count", 0)
            ctx.setdefault("endorsed_to_them", "0.00")
            ctx.setdefault("endorsed_to_them_count", 0)
            ctx.setdefault("mutual_guarantee", "0.00")
            ctx.setdefault("contract_no", "—")
            ctx.setdefault("contract_period", "—")
            ctx.setdefault("payment_terms", "—")
            ctx.setdefault("credit_limit", 0.0)
            ctx.setdefault("title_retention", "无")
            ctx.setdefault("dispute_resolution", "诉讼至本公司所在地法院")
            ctx.setdefault("contract_amendments", "无")
            ctx.setdefault("auditor_zip", "—")
            ctx.setdefault("litigation_status", "无")
            return CUSTOMER_SUPPLIER_LETTER_TEMPLATE.format(**ctx)
        if template_id == "other_std":
            return OTHER_RECEIVABLE_PAYABLE_TEMPLATE.format(**ctx)

        raise LetterGenerationError(f"未知模板: {template_id}")

    # ---- 渲染为 docx -------------------------------------------------

    def render_docx(
        self,
        template_id: str,
        text: str,
        *,
        filename_hint: str,
        meta: dict[str, Any],
    ) -> Path:
        """将纯文本函证转为 docx（用 python-docx）。

        P0 修复: filename_hint 保留中文字符, 仅替换文件系统不接受的特殊字符, 并加 uuid 防覆盖.
        """
        try:
            from docx import Document  # type: ignore
            from docx.shared import Pt
        except ImportError as exc:
            raise LetterGenerationError(
                "生成 docx 需要安装 python-docx：`uv add python-docx`"
            ) from exc

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "SimSun"
        style.font.size = Pt(10.5)

        # 标题按 template_id 区分 (P0 修复)
        title_map = {
            "bank_official": f"银行询证函 - {meta.get('party_name', '')}",
            "customer_std": f"应收账款询证函 - {meta.get('party_name', '')}",
            "supplier_std": f"应付账款询证函 - {meta.get('party_name', '')}",
            "other_std": f"其他往来询证函 - {meta.get('party_name', '')}",
        }
        title = doc.add_heading(
            title_map.get(template_id, f"询证函 - {meta.get('party_name', '')}"), level=1
        )
        title.alignment = 1

        for line in text.splitlines():
            if not line.strip():
                doc.add_paragraph("")
                continue
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(10.5)

        # 文件名: 替换系统不接受的字符, 保留中文; 加 uuid 防覆盖
        import uuid

        safe = re.sub(r"[\x00-\x1f/\\:*?\"<>|]", "_", filename_hint)[:80]
        suffix = uuid.uuid4().hex[:6]
        path = self.output_dir / f"{safe}_{suffix}.docx"
        doc.save(str(path))
        return path

    # ---- 渲染为 PDF (可选，需 docx2pdf / libreoffice) -----------------

    def render_pdf(self, docx_path: Path) -> Optional[Path]:
        """可选 PDF 转换（依赖 libreoffice / docx2pdf）。失败返回 None。"""
        try:
            import subprocess

            out = self.output_dir / (docx_path.stem + ".pdf")
            r = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(self.output_dir),
                    str(docx_path),
                ],
                capture_output=True,
                timeout=60,
            )
            if r.returncode == 0 and out.exists():
                return out
        except Exception:  # noqa: BLE001
            pass
        return None

    # ---- 一键生成 ----------------------------------------------------

    def generate(
        self,
        template_id: str,
        *,
        company_name: str,
        period: str,
        period_start: str,
        balance_date: str,
        sent_date: str,
        recipient: str,
        cpa_firm: str = "××会计师事务所(特殊普通合伙)",
        cpa_address: str = "北京市××区××路×号",
        auditor_name: str = "审计师",
        auditor_phone: str = "010-××××××××",
        party_name: str = "",
        file_format: str = "docx",
        **kwargs: Any,
    ) -> tuple[Path, str, str]:
        """生成询证函文件 + 返回 (path, content_text, actual_file_format).

        P0 修复: 返回值新增 actual_file_format — 当用户请求 pdf 但转换失败时
        落库的 file_format 应该是 docx, 不能是 pdf.
        """
        text = self.render_text(
            template_id,
            company_name=company_name,
            period=period,
            period_start=period_start,
            balance_date=balance_date,
            sent_date=sent_date,
            recipient=recipient,
            cpa_firm=cpa_firm,
            cpa_address=cpa_address,
            auditor_name=auditor_name,
            auditor_phone=auditor_phone,
            **kwargs,
        )
        meta = {"party_name": party_name, "template_id": template_id}
        if file_format == "pdf":
            docx_path = self.render_docx(
                template_id,
                text,
                filename_hint=f"{template_id}_{party_name}_{sent_date}",
                meta=meta,
            )
            pdf = self.render_pdf(docx_path)
            if pdf is not None and pdf.exists():
                return pdf, text, "pdf"
            # P0 修复: PDF 转换失败, 落库 file_format 应为 docx
            return docx_path, text, "docx"
        path = self.render_docx(
            template_id,
            text,
            filename_hint=f"{template_id}_{party_name}_{sent_date}",
            meta=meta,
        )
        return path, text, "docx"
