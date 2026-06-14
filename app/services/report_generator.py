"""综合报告生成器 - 第六阶段."""

from typing import Dict, List, Optional
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ComprehensiveReportGenerator:
    """生成综合分析报告."""

    def __init__(self):
        self.report_date = datetime.now().strftime("%Y年%m月%d日")

    def generate_word_report(
        self,
        project_info: Dict,
        financial_data: Dict,
        risk_analysis: Dict,
        trial_balance: Dict,
        regulatory_cases: List[Dict],
    ) -> bytes:
        """生成Word格式综合报告."""
        doc = Document()

        # 标题
        title = doc.add_heading(f"{project_info.get('company_name', '')} IPO审计综合分析报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 报告信息
        doc.add_paragraph(f"报告日期：{self.report_date}")
        doc.add_paragraph(f"项目名称：{project_info.get('name', '')}")
        doc.add_paragraph(f"审计年度：{project_info.get('fiscal_year', '')}年")
        doc.add_paragraph("")

        # 1. 企业概况
        doc.add_heading("一、企业概况", 1)
        self._add_company_overview(doc, project_info, financial_data)

        # 2. 财务分析
        doc.add_heading("二、财务指标分析", 1)
        self._add_financial_analysis(doc, financial_data)

        # 3. 风险分析
        doc.add_heading("三、风险分析", 1)
        self._add_risk_analysis(doc, risk_analysis)

        # 4. 监管关注点
        doc.add_heading("四、监管关注点", 1)
        self._add_regulatory_focus(doc, regulatory_cases)

        # 5. 审计结论
        doc.add_heading("五、审计结论", 1)
        self._add_audit_conclusion(doc, risk_analysis)

        # 保存
        from io import BytesIO

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _add_company_overview(self, doc: Document, project_info: Dict, financial_data: Dict):
        """添加企业概况."""
        table = doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"

        rows_data = [
            ("公司名称", project_info.get("company_name", "")),
            ("所属行业", project_info.get("industry", "")),
            ("注册资本", f"{financial_data.get('registered_capital') or 0:,.2f}万元"),
            ("总资产", f"{financial_data.get('total_assets') or 0:,.2f}万元"),
            ("营业收入", f"{financial_data.get('revenue') or 0:,.2f}万元"),
        ]

        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        doc.add_paragraph("")

    def _add_financial_analysis(self, doc: Document, financial_data: Dict):
        """添加财务分析."""
        # 盈利能力
        doc.add_heading("1. 盈利能力", 2)
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        rows_data = [
            ("毛利率", f"{financial_data.get('gross_margin') or 0:.2f}%"),
            ("净利率", f"{financial_data.get('net_margin') or 0:.2f}%"),
            ("净资产收益率", f"{financial_data.get('roe') or 0:.2f}%"),
            ("每股收益", f"{financial_data.get('eps') or 0:.2f}元"),
        ]

        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        # 偿债能力
        doc.add_heading("2. 偿债能力", 2)
        table2 = doc.add_table(rows=3, cols=2)
        table2.style = "Table Grid"

        rows_data2 = [
            ("流动比率", f"{financial_data.get('current_ratio') or 0:.2f}"),
            ("速动比率", f"{financial_data.get('quick_ratio') or 0:.2f}"),
            ("资产负债率", f"{financial_data.get('debt_ratio') or 0:.2f}%"),
        ]

        for i, (label, value) in enumerate(rows_data2):
            table2.rows[i].cells[0].text = label
            table2.rows[i].cells[1].text = str(value)

        doc.add_paragraph("")

    def _add_risk_analysis(self, doc: Document, risk_analysis: Dict):
        """添加风险分析."""
        risk_level = risk_analysis.get("risk_level", "中")
        risk_points = risk_analysis.get("risk_points", [])

        doc.add_paragraph(f"综合风险等级：{risk_level}")
        doc.add_paragraph("")

        if risk_points:
            doc.add_heading("主要风险点：", 2)
            for i, point in enumerate(risk_points, 1):
                doc.add_paragraph(f"{i}. {point}", style="List Bullet")
        else:
            doc.add_paragraph("未发现重大风险点。")

        doc.add_paragraph("")

    def _add_regulatory_focus(self, doc: Document, regulatory_cases: List[Dict]):
        """添加监管关注点."""
        if not regulatory_cases:
            doc.add_paragraph("未找到相关监管案例。")
            return

        doc.add_paragraph(f"找到{len(regulatory_cases)}条相关监管案例：")
        doc.add_paragraph("")

        for i, case in enumerate(regulatory_cases[:5], 1):
            doc.add_heading(f"案例{i}：{case.get('title', '')}", 3)
            doc.add_paragraph(f"来源：{case.get('source', '')}")
            doc.add_paragraph(f"日期：{case.get('publish_date', '')}")
            doc.add_paragraph(f"内容摘要：{case.get('content', '')[:200]}...")
            doc.add_paragraph("")

    def _add_audit_conclusion(self, doc: Document, risk_analysis: Dict):
        """添加审计结论."""
        risk_level = risk_analysis.get("risk_level", "中")

        conclusions = {
            "高": "经分析，该公司存在较高风险，建议加强审计程序，重点关注风险事项。",
            "中": "经分析，该公司风险适中，建议保持常规审计程序。",
            "低": "经分析，该公司风险较低，可以适当简化审计程序。",
        }

        doc.add_paragraph(conclusions.get(risk_level, ""))
        doc.add_paragraph("")
        doc.add_paragraph("审计人员：________________")
        doc.add_paragraph("")
        doc.add_paragraph("复核人员：________________")


class InteractiveReportGenerator:
    """生成交互式Web报告数据."""

    @staticmethod
    def generate_dashboard_data(
        project_info: Dict,
        financial_data: Dict,
        risk_analysis: Dict,
        trial_balance: Dict,
        anomalies: List[Dict],
    ) -> Dict:
        """生成仪表盘数据."""
        return {
            "project": {
                "name": project_info.get("name", ""),
                "company": project_info.get("company_name", ""),
                "industry": project_info.get("industry", ""),
                "fiscal_year": project_info.get("fiscal_year", ""),
            },
            "financial_summary": {
                "total_assets": financial_data.get("total_assets", 0),
                "revenue": financial_data.get("revenue", 0),
                "net_profit": financial_data.get("net_profit", 0),
                "gross_margin": financial_data.get("gross_margin", 0),
            },
            "risk_assessment": {
                "level": risk_analysis.get("risk_level", "中"),
                "risk_points": risk_analysis.get("risk_points", []),
                "recommendations": risk_analysis.get("recommendations", []),
            },
            "balance_status": {
                "is_balanced": trial_balance.get("is_balanced", False),
                "total_debit": trial_balance.get("total_debit", 0),
                "total_credit": trial_balance.get("total_credit", 0),
                "difference": trial_balance.get("difference", 0),
            },
            "anomalies": anomalies,
            "charts": {
                "asset_structure": {
                    "labels": ["流动资产", "非流动资产"],
                    "values": [
                        financial_data.get("current_assets", 0),
                        financial_data.get("non_current_assets", 0),
                    ],
                },
                "risk_heatmap": {
                    "categories": ["收入确认", "应收账款", "存货", "关联交易", "商誉"],
                    "scores": [
                        risk_analysis.get("scores", {}).get(cat, 50)
                        for cat in ["收入确认", "应收账款", "存货", "关联交易", "商誉"]
                    ],
                },
            },
        }

    @staticmethod
    def generate_trend_data(financial_data_historical: List[Dict]) -> Dict:
        """生成趋势数据."""
        return {
            "revenue_trend": [
                {"year": d.get("year", ""), "value": d.get("revenue", 0)}
                for d in financial_data_historical
            ],
            "profit_trend": [
                {"year": d.get("year", ""), "value": d.get("net_profit", 0)}
                for d in financial_data_historical
            ],
            "margin_trend": [
                {"year": d.get("year", ""), "value": d.get("gross_margin", 0)}
                for d in financial_data_historical
            ],
        }


class PDFReportGenerator:
    """生成PDF格式报告."""

    @staticmethod
    def generate_pdf(
        project_info: Dict,
        financial_data: Dict,
        risk_analysis: Dict,
        trial_balance: Dict,
    ) -> bytes:
        """生成PDF报告."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import cm

        from io import BytesIO

        buffer = BytesIO()

        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # 标题
        title_style = ParagraphStyle(
            "CustomTitle",
            parent=styles["Heading1"],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # CENTER
        )
        story.append(
            Paragraph(f"{project_info.get('company_name', '')} IPO审计综合分析报告", title_style)
        )
        story.append(Spacer(1, 0.5 * cm))

        # 基本信息
        story.append(Paragraph("一、基本信息", styles["Heading2"]))
        info_data = [
            ["项目名称", project_info.get("name", "")],
            ["公司名称", project_info.get("company_name", "")],
            ["所属行业", project_info.get("industry", "")],
            ["审计年度", f"{project_info.get('fiscal_year', '')}年"],
        ]
        info_table = Table(info_data, colWidths=[4 * cm, 10 * cm])
        info_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ]
            )
        )
        story.append(info_table)
        story.append(Spacer(1, 0.5 * cm))

        # 财务指标
        story.append(Paragraph("二、主要财务指标", styles["Heading2"]))
        fin_data = [
            ["指标", "金额/比例"],
            ["总资产", f"{financial_data.get('total_assets') or 0:,.2f}万元"],
            ["营业收入", f"{financial_data.get('revenue') or 0:,.2f}万元"],
            ["净利润", f"{financial_data.get('net_profit') or 0:,.2f}万元"],
            ["毛利率", f"{financial_data.get('gross_margin') or 0:.2f}%"],
        ]
        fin_table = Table(fin_data, colWidths=[4 * cm, 10 * cm])
        fin_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                ]
            )
        )
        story.append(fin_table)
        story.append(Spacer(1, 0.5 * cm))

        # 风险评估
        story.append(Paragraph("三、风险评估", styles["Heading2"]))
        story.append(
            Paragraph(f"综合风险等级：{risk_analysis.get('risk_level', '中')}", styles["Normal"])
        )
        story.append(Spacer(1, 0.3 * cm))

        risk_points = risk_analysis.get("risk_points", [])
        if risk_points:
            story.append(Paragraph("主要风险点：", styles["Normal"]))
            for i, point in enumerate(risk_points, 1):
                story.append(Paragraph(f"{i}. {point}", styles["Normal"]))
        story.append(Spacer(1, 0.5 * cm))

        # 试算平衡
        story.append(Paragraph("四、试算平衡", styles["Heading2"]))
        story.append(
            Paragraph(
                f"状态：{'平衡' if trial_balance.get('is_balanced') else '不平衡'} | "
                f"借方合计：{trial_balance.get('total_debit') or 0:,.2f} | "
                f"贷方合计：{trial_balance.get('total_credit') or 0:,.2f}",
                styles["Normal"],
            )
        )

        doc.build(story)
        buffer.seek(0)
        return buffer.read()


class ReportScheduler:
    """报告调度器 - 支持定期生成报告."""

    def __init__(self):
        self.scheduled_reports: List[Dict] = []

    def schedule_report(
        self,
        project_id: int,
        report_type: str,
        frequency: str,  # daily, weekly, monthly
        recipients: List[str],
    ) -> Dict:
        """调度报告生成."""
        schedule = {
            "id": len(self.scheduled_reports) + 1,
            "project_id": project_id,
            "report_type": report_type,
            "frequency": frequency,
            "recipients": recipients,
            "created_at": datetime.now().isoformat(),
            "status": "active",
        }
        self.scheduled_reports.append(schedule)
        return schedule

    def get_scheduled_reports(self, project_id: Optional[int] = None) -> List[Dict]:
        """获取调度的报告列表."""
        if project_id:
            return [r for r in self.scheduled_reports if r["project_id"] == project_id]
        return self.scheduled_reports

    def cancel_schedule(self, schedule_id: int) -> bool:
        """取消报告调度."""
        for schedule in self.scheduled_reports:
            if schedule["id"] == schedule_id:
                schedule["status"] = "cancelled"
                return True
        return False
