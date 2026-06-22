"""Markdown → .docx — 照搬 app.services.report_generator.ComprehensiveReportGenerator 风格.

差异:
- 接受 markdown 字符串, 简单逐段解析
- 落盘路径: settings.SENTIMENT_OUTPUT_DIR / "briefings" / "{project_id}_{date}.docx"
- 返回 (file_path, sha256)
"""

from __future__ import annotations

import hashlib
import logging
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt

from app.core.config import settings

logger = logging.getLogger(__name__)

# 固定时间戳 — 同一份 markdown 导出, SHA256 应稳定.
# 不用 ``datetime.now()``: python-docx ``CorePropertiesPart.default`` 会在
# ``Document()`` 构造时把 ``modified`` 设为当前时间, 导致两次 save
# 出 .docx 内 docProps/core.xml 的 <dcterms:created>/<dcterms:modified>
# 不一致 → SHA256 漂移 → 同内容 hash 不同. 固化 sentinel 后, 同内容
# hash 稳定, 既满足"幂等导出"业务承诺, 也让回归测试可比对.
STABLE_DT = datetime(2024, 1, 1, 0, 0, 0, tzinfo=timezone.utc)

# ``docProps/core.xml`` 里的两个时间戳元素名, 用 regex 原地替换.
_CORE_CREATED_RE = re.compile(
    rb'(<dcterms:created[^>]*>)[^<]*(</dcterms:created>)'
)
_CORE_MODIFIED_RE = re.compile(
    rb'(<dcterms:modified[^>]*>)[^<]*(</dcterms:modified>)'
)
_STABLE_TS_BYTES = STABLE_DT.strftime("%Y-%m-%dT%H:%M:%SZ").encode("ascii")


def _strip_docx_timestamps(docx_bytes: bytes) -> bytes:
    """把 docx (zip) 内的 docProps/core.xml 时间戳替换为 sentinel.

    ``zipfile`` 重新打包, 把 ``docProps/core.xml`` 的 <dcterms:created> /
    <dcterms:modified> 替换为 STABLE_DT; 其他 entry 原样返回.
    这样同 markdown 内容两次 hash 完全相同.
    """
    import io

    src = io.BytesIO(docx_bytes)
    out = io.BytesIO()
    with zipfile.ZipFile(src, "r") as zin, zipfile.ZipFile(
        out, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "docProps/core.xml":
                data = _CORE_CREATED_RE.sub(
                    rb"\1" + _STABLE_TS_BYTES + rb"\2", data
                )
                data = _CORE_MODIFIED_RE.sub(
                    rb"\1" + _STABLE_TS_BYTES + rb"\2", data
                )
            # 沿用原 ZipInfo (compress_type/date_time), 只覆盖内容.
            zout.writestr(item, data)
    return out.getvalue()


def _stable_docx_sha256(path: Path) -> str:
    """读 .docx, 抹平 core.xml 时间戳, 算 SHA256.

    同 markdown 多次导出, SHA 必相同; 任何 markdown 变化, SHA 必不同.
    """
    raw = Path(path).read_bytes()
    stable = _strip_docx_timestamps(raw)
    return hashlib.sha256(stable).hexdigest()


class BriefingWordExporter:
    """简报 Markdown → .docx.

    极简实现: 把 Markdown 按行处理
        - # / ## / ### → heading
        - | col | col |  → table (简单支持)
        - 1. / 2.       → numbered list
        - - item        → bullet
        - 其它          → paragraph
    不追求完整 Markdown 渲染, 但能稳定呈现 4 轮 LLM 的产物.
    """

    def __init__(self) -> None:
        self.output_dir = settings.SENTIMENT_OUTPUT_DIR / "briefings"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export(
        self,
        project_id: int,
        briefing_date: str,
        company_name: str,
        markdown: str,
        *,
        version_no: Optional[int] = None,
    ) -> tuple[Path, str]:
        """生成 .docx, 落盘, 返回 (path, sha256)."""
        doc = Document()
        # 默认中文字体
        style = doc.styles["Normal"]
        style.font.size = Pt(10.5)
        style.font.name = "Microsoft YaHei"

        # 顶部标题
        doc.add_heading(f"{company_name} {briefing_date} 舆情简报", level=0)

        self._render_markdown_into_doc(doc, markdown)

        # 落盘
        suffix = f"_v{version_no}" if version_no and version_no > 1 else ""
        fname = f"project_{project_id}_{briefing_date}{suffix}.docx"
        path = self.output_dir / fname
        doc.save(str(path))

        # sha256
        digest = _stable_docx_sha256(path)
        logger.info("BriefingWordExporter: %s -> %s (sha256=%s)", company_name, path, digest[:12])
        return path, digest

    def _render_markdown_into_doc(self, doc: Document, markdown: str) -> None:
        lines = markdown.splitlines()
        i = 0
        in_table = False
        table_buf: list[list[str]] = []

        def flush_table() -> None:
            nonlocal table_buf, in_table
            if not table_buf:
                in_table = False
                return
            cols = max(len(r) for r in table_buf)
            t = doc.add_table(rows=len(table_buf), cols=cols)
            t.style = "Table Grid"
            for r_idx, row in enumerate(table_buf):
                for c_idx in range(cols):
                    val = row[c_idx] if c_idx < len(row) else ""
                    t.cell(r_idx, c_idx).text = val
            table_buf = []
            in_table = False
            doc.add_paragraph("")  # 表后空行

        while i < len(lines):
            line = lines[i].rstrip()
            stripped = line.strip()

            # 表格行
            if stripped.startswith("|") and stripped.endswith("|"):
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                # 跳过分隔行 | --- | --- |
                if all(re.match(r"^[-:\s]+$", c) for c in cells):
                    i += 1
                    continue
                table_buf.append(cells)
                in_table = True
                i += 1
                continue
            else:
                if in_table:
                    flush_table()

            # 标题
            m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if m:
                level = min(len(m.group(1)), 9)
                doc.add_heading(m.group(2), level=level)
                i += 1
                continue

            # 引用块
            if stripped.startswith(">"):
                doc.add_paragraph(stripped.lstrip(">").strip(), style="Intense Quote")
                i += 1
                continue

            # 数字列表
            m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
            if m:
                doc.add_paragraph(f"{m.group(1)}. {m.group(2)}", style="List Number")
                i += 1
                continue

            # 无序列表
            if stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
                i += 1
                continue

            # 空行
            if not stripped:
                doc.add_paragraph("")
                i += 1
                continue

            # 普通段落
            doc.add_paragraph(stripped)
            i += 1

        if in_table:
            flush_table()
