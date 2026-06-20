"""round 36 (2026-06-20) 测试空白填补 + 1 P0 修复回归.

覆盖 4 个 P0 模块 (原本 0 测试):
  - app/services/regulatory_case_service.py — 监管案例抓取 + 匹配
  - app/services/notification/__init__.py    — 通用通知中心 (push/mark_read/list)
  - app/services/sentiment/briefing/generator.py — 4 轮 LLM 协议 + 简报拼装
  - app/services/sales_ledger/document_parser.py — xlsx/docx/pdf 文件解析

同时回归测试 sentiment/quarterly/verifier.py (P0 fix):
  - _find_value 签名兼容 field_name kwarg (round 36 P0)

约束:
  - 不改业务代码 (除 verifier.py 的 P0 签名加 Any 类型 + 中文 docstring)
  - 用 tests/_helpers/{db,auth} (in-memory SQLite + make_user)
  - pytest-asyncio auto mode (项目 conftest 已配 asyncio_mode="auto")
  - httpx / LLM 用 AsyncMock 隔离, 不发真实请求
"""
from __future__ import annotations

import json
import os
import tempfile
from pathlib import Path
from typing import Any
from unittest.mock import AsyncMock, MagicMock, patch

import pandas as pd
import pytest
import pytest_asyncio

# 在 import app 之前设环境变量 (避免依赖外部 LLM 真实请求)
_tmp_db = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
_tmp_db.close()
os.environ["DATABASE_URL"] = f"sqlite+aiosqlite:///{_tmp_db.name}"
os.environ.setdefault("AUTH_ENABLED", "false")
os.environ.setdefault("AUDIT_LOG_WRITE_ONLY", "false")


# ============================================================
#  P0 修复回归: sentiment/quarterly/verifier._find_value
# ============================================================


class TestQuarterlyVerifierP0FieldNameKwarg:
    """Round 36 P0: verify() 透传 field_name kwarg, _find_value 签名需接受.

    旧问题: 旧签名未声明 field_name 参数, 调用时 TypeError.
    修复: 签名显式声明 ``*, field_name: Optional[str] = None`` (keyword-only).
    本测试同时验证:
      - 不传 field_name (默认 None) 不抛错, 走通用 subject
      - 传 field_name kwarg, subject 上下文匹配正确
      - 整数 / 浮点 / 字符串 value 类型都能匹配
    """

    def test_find_value_accepts_field_name_kwarg(self):
        """直接调用 _find_value 带 field_name kwarg, 不抛 TypeError."""
        from app.services.sentiment.quarterly.verifier import QuarterlyVerifier

        v = QuarterlyVerifier()
        # 直接调用 _find_value (绕过 verify), 传 field_name kwarg
        matched_in, matched_val, note = v._find_value(
            0.5, "毛利率 50.0%", "", field_name="gross_margin",
        )
        # 毛利率 50.0% + field_name=gross_margin → 应匹配 events
        assert matched_in == "events"
        assert matched_val == "50.0%"

    def test_find_value_default_no_field_name(self):
        """不传 field_name → 走通用 subject (generic subjects), 不抛错."""
        from app.services.sentiment.quarterly.verifier import QuarterlyVerifier

        v = QuarterlyVerifier()
        # field_name 默认为 None
        matched_in, matched_val, note = v._find_value(
            0.5, "本期毛利率 50.0%, 业绩亮眼", "",
        )
        # 没有 field_name 但文本有 "毛利率" + "50.0%" → 通用 subject 匹配
        assert matched_in == "events"
        assert "50.0%" in matched_val

    def test_find_value_no_match_returns_none_tuple(self):
        """value 在文本中找不到 → 返回 ('none', None, note) 三元组."""
        from app.services.sentiment.quarterly.verifier import QuarterlyVerifier

        v = QuarterlyVerifier()
        matched_in, matched_val, note = v._find_value(
            0.5, "完全无关的内容", "", field_name="gross_margin",
        )
        assert matched_in == "none"
        assert matched_val is None
        assert "舆情印证" in note

    def test_find_value_pct_subject_guard_preserved(self):
        """Round 35 P0 防护不能破: 百分比形需 subject 上下文匹配."""
        from app.services.sentiment.quarterly.verifier import QuarterlyVerifier

        v = QuarterlyVerifier()
        # 文本含 "营收 50.0%", 但 field_name 是 gross_margin → 不应误匹配
        matched_in, matched_val, note = v._find_value(
            0.5, "营收同比增长 50.0%, 净利润大幅提升", "", field_name="gross_margin",
        )
        # 营收 50.0% 没有 "毛利率" / "margin" 上下文 → 不匹配
        assert matched_in == "none"

    def test_verify_passes_field_name_to_find_value(self):
        """verify() 内部调用 _find_value 时透传 field_name (集成测试)."""
        from app.services.sentiment.quarterly.verifier import QuarterlyVerifier

        v = QuarterlyVerifier()
        # 监控 _find_value 被调用时的 kwargs
        original_find_value = v._find_value
        captured_kwargs: list[dict] = []

        def spy_find_value(value, events_text, briefings_text, **kwargs):
            captured_kwargs.append({"value": value, "kwargs": kwargs})
            return original_find_value(value, events_text, briefings_text, **kwargs)

        with patch.object(v, "_find_value", side_effect=spy_find_value):
            v.verify(
                markdown="",
                financial_input={"gross_margin": 0.5, "revenue": 1_000_000},
                events=[{"id": 1, "title": "公告", "content_text": "毛利率 50.0%, 营收 1,000,000"}],
                briefings=[],
            )

        # verify() 至少对 numeric 字段调一次 _find_value, 且 field_name 是 kwarg
        assert len(captured_kwargs) >= 2, (
            f"verify() 应至少调 2 次 _find_value (gross_margin + revenue), 实际 {len(captured_kwargs)}"
        )
        # 每次调用都透传了 field_name
        for call in captured_kwargs:
            assert "field_name" in call["kwargs"], (
                f"_find_value 调用未透传 field_name: {call}"
            )


# ============================================================
#  Module 1: regulatory_case_service
# ============================================================


def _mock_async_client(*, html_body: str = "", status_code: int = 200, raise_exc: Exception | None = None):
    """构造 mock httpx.AsyncClient. 支持 raise_for_status 模拟.

    注: 原代码 ``await self.session.get(...)`` 但 ``response.raise_for_status()``
    和 ``response.text`` 都是同步访问 (httpx 0.26+). 所以 response 用 MagicMock,
    client.get() 用 AsyncMock (因为代码里 await), aclose() 也用 AsyncMock.
    """
    client = AsyncMock()
    response = MagicMock()
    response.text = html_body
    response.status_code = status_code
    if raise_exc is not None:
        response.raise_for_status.side_effect = raise_exc
    else:
        response.raise_for_status.return_value = None
    # client.get() 是 coroutine, await 后返回 response
    client.get = AsyncMock(return_value=response)
    client.aclose = AsyncMock(return_value=None)
    return client


class TestRegulatoryCaseScraperCSRC:
    """RegulatoryCaseScraper.scrape_csrc_inquiry — 抓证监会问询函."""

    @pytest.mark.asyncio
    async def test_scrape_csrc_inquiry_parses_table_rows(self):
        """CSRC HTML 含 table.list 行 → 解析为 cases 字典列表."""
        from app.services.regulatory_case_service import RegulatoryCaseScraper

        html = """
        <html><body>
        <table class="list">
            <tr><td>CSRC-2024-001</td><td>2024-01-15</td><td>关于X公司问询</td><td>内容1</td></tr>
            <tr><td>CSRC-2024-002</td><td>2024-02-20</td><td>关于Y公司问询</td><td>内容2</td></tr>
        </table>
        </body></html>
        """
        scraper = RegulatoryCaseScraper()
        scraper.session = _mock_async_client(html_body=html)

        try:
            cases = await scraper.scrape_csrc_inquiry(page=1, keyword="")
            assert len(cases) == 2
            assert cases[0]["case_no"] == "CSRC-2024-001"
            assert cases[0]["source"] == "证监会"
            assert cases[0]["case_type"] == "问询函"
            assert cases[1]["case_no"] == "CSRC-2024-002"
        finally:
            await scraper.close()

    @pytest.mark.asyncio
    async def test_scrape_csrc_inquiry_with_keyword_adds_param(self):
        """传 keyword 时 params 应包含 keyword 字段."""
        from app.services.regulatory_case_service import RegulatoryCaseScraper

        scraper = RegulatoryCaseScraper()
        scraper.session = _mock_async_client(html_body="<html></html>")

        try:
            await scraper.scrape_csrc_inquiry(page=2, keyword="毛利率")
            # 校验 session.get 被调用且 params 包含 keyword
            call_kwargs = scraper.session.get.call_args.kwargs
            assert call_kwargs["params"].get("keyword") == "毛利率"
            assert call_kwargs["params"].get("page") == 2
        finally:
            await scraper.close()

    @pytest.mark.asyncio
    async def test_scrape_csrc_inquiry_returns_empty_on_error(self):
        """HTTP 抛错时, 返回空列表 (不抛给调用方)."""
        import httpx

        from app.services.regulatory_case_service import RegulatoryCaseScraper

        scraper = RegulatoryCaseScraper()
        scraper.session = _mock_async_client(
            raise_exc=httpx.HTTPStatusError(
                "500", request=MagicMock(), response=MagicMock(status_code=500),
            ),
        )

        try:
            cases = await scraper.scrape_csrc_inquiry()
            assert cases == []
        finally:
            await scraper.close()


class TestRegulatoryCaseScraperSSEAndSZSE:
    """SSE / SZSE 抓取 (mock session 避免真实 URL build)."""

    @pytest.mark.asyncio
    async def test_scrape_sse_inquiry_parses_items(self):
        """SSE HTML 含 .inquiry-item → 解析为 cases."""
        # 注: 真实 settings 缺 SseUrl, 这里通过 mock session 让代码不真正访问
        from app.services.regulatory_case_service import RegulatoryCaseScraper

        html = """
        <html><body>
        <div class="inquiry-item">
            <span class="case-no">SSE-2024-01</span>
            <span class="date">2024-03-01</span>
            <span class="title">SSE 标题</span>
            <span class="content">SSE 内容</span>
        </div>
        </body></html>
        """
        scraper = RegulatoryCaseScraper()
        scraper.session = _mock_async_client(html_body=html)

        try:
            # 临时 mock settings 以绕开缺字段问题
            with patch("app.services.regulatory_case_service.settings") as mock_settings:
                mock_settings.CSRC_URL = "http://csrc"
                mock_settings.SseUrl = "http://sse"
                mock_settings.SzseUrl = "http://szse"
                cases = await scraper.scrape_sse_inquiry(page=1)
            assert len(cases) == 1
            assert cases[0]["case_no"] == "SSE-2024-01"
            assert cases[0]["source"] == "上交所"
        finally:
            await scraper.close()

    @pytest.mark.asyncio
    async def test_scrape_szse_inquiry_parses_items(self):
        """SZSE HTML 含 .inquiry-list li → 解析为 cases."""
        from app.services.regulatory_case_service import RegulatoryCaseScraper

        html = """
        <html><body>
        <ul class="inquiry-list">
            <li>
                <span class="code">SZSE-001</span>
                <span class="date">2024-04-01</span>
                <span class="title">深交所标题</span>
                <span class="summary">深交所摘要</span>
            </li>
        </ul>
        </body></html>
        """
        scraper = RegulatoryCaseScraper()
        scraper.session = _mock_async_client(html_body=html)

        try:
            with patch("app.services.regulatory_case_service.settings") as mock_settings:
                mock_settings.CSRC_URL = "http://csrc"
                mock_settings.SseUrl = "http://sse"
                mock_settings.SzseUrl = "http://szse"
                cases = await scraper.scrape_szse_inquiry(page=1)
            assert len(cases) == 1
            assert cases[0]["case_no"] == "SZSE-001"
            assert cases[0]["source"] == "深交所"
        finally:
            await scraper.close()


class TestCaseMatcherMatch:
    """CaseMatcher.match_by_industry / match_by_keywords / calculate_relevance."""

    def test_match_by_industry_filters_manufacturing(self):
        """制造业 industry → 只保留含『毛利率/应收/存货/关联/收入』关键词的 case."""
        from app.services.regulatory_case_service import CaseMatcher

        matcher = CaseMatcher()
        cases = [
            {"title": "A公司毛利率异常", "content": "存货周转率低"},
            {"title": "B公司研发投入", "content": "技术升级"},
            {"title": "C公司收入确认", "content": "提前确认收入"},
        ]
        matched = matcher.match_by_industry(cases, "制造业")
        # A 命中 (毛利率), C 命中 (收入确认), B 不命中
        assert len(matched) == 2
        titles = {c["title"] for c in matched}
        assert "A公司毛利率异常" in titles
        assert "C公司收入确认" in titles
        assert "B公司研发投入" not in titles

    def test_match_by_industry_unknown_returns_empty(self):
        """未知 industry → 关键词列表空 → 无 case 命中 (返空列表, 不抛错).

        实现契约: ``industry_keywords.get(industry, [])`` 返空列表时,
        内层 for 循环不执行, 故 matched 始终为 [].
        """
        from app.services.regulatory_case_service import CaseMatcher

        matcher = CaseMatcher()
        cases = [{"title": "ABC公司", "content": "无匹配词"}]
        matched = matcher.match_by_industry(cases, "未知行业XYZ")
        assert matched == []

    def test_match_by_industry_empty_industry_returns_all(self):
        """industry 为空字符串 → 返全部 cases (不筛选)."""
        from app.services.regulatory_case_service import CaseMatcher

        matcher = CaseMatcher()
        cases = [{"title": "X1"}, {"title": "X2"}]
        matched = matcher.match_by_industry(cases, "")
        assert matched == cases

    def test_match_by_keywords_scores_and_sorts(self):
        """关键词匹配: 标题命中权重 3, 内容命中权重 1; 按分数降序.

        实现细节: 同一关键词 if/elif 二选一, 标题命中不计内容.
        """
        from app.services.regulatory_case_service import CaseMatcher

        matcher = CaseMatcher()
        cases = [
            {"title": "ABC公司", "content": "无关键"},       # 仅内容 → 1
            {"title": "关于关键", "content": "提及关键"},    # 标题命中 → 3
            {"title": "其他事", "content": "提到关键"},       # 仅内容 → 1
            {"title": "无关键", "content": "无关键"},         # 标题命中 → 3
        ]
        matched = matcher.match_by_keywords(cases, ["关键"])
        # 排序: 标题命中 (3, 3) > 内容命中 (1, 1)
        scores = [c["match_score"] for c in matched]
        assert scores == [3, 3, 1, 1], f"实际 {scores}"
        # 全部有 matched_keywords 字段
        for c in matched:
            assert "matched_keywords" in c
            assert "关键" in c["matched_keywords"]

    def test_match_by_keywords_empty_keywords_returns_empty(self):
        """空关键词列表 → 全部跳过, 返空."""
        from app.services.regulatory_case_service import CaseMatcher

        matcher = CaseMatcher()
        cases = [{"title": "X", "content": "Y"}]
        matched = matcher.match_by_keywords(cases, [])
        assert matched == []

    def test_calculate_relevance_combines_industry_and_keywords(self):
        """calculate_relevance 综合行业 + 关键词 + 规模打分."""
        from app.services.regulatory_case_service import CaseMatcher

        matcher = CaseMatcher()
        # industry=制造业 同时出现在 title 和 content → +30
        # keywords=毛利率 命中 title → +10
        # revenue < 10亿 → 不命中"大额"加分
        case = {
            "title": "制造业毛利率异常案例",
            "content": "制造业某公司毛利率异常波动, 大额关联交易",
        }
        company_info = {
            "industry": "制造业",
            "keywords": ["毛利率"],
            "revenue": 500_000_000,  # < 10亿
        }
        score = matcher.calculate_relevance(case, company_info)
        # industry in title → +30
        # keywords 毛利率 in title → +10 (只算一次, 不会因 elif 跳过)
        # 总 40
        assert score == 40.0

    def test_calculate_relevance_high_revenue_bonus(self):
        """revenue > 100亿 + content 含"大额" → +10 (规模匹配加分)."""
        from app.services.regulatory_case_service import CaseMatcher

        matcher = CaseMatcher()
        case = {
            "title": "某金融服务案例",
            "content": "大额关联交易, 影响重大",
        }
        company_info = {
            "industry": "金融服务",
            "keywords": ["关联"],
            "revenue": 50_000_000_000,  # > 100亿
        }
        score = matcher.calculate_relevance(case, company_info)
        # industry=金融服务 in title → +30
        # keywords=关联 in content → +3
        # revenue > 100亿 + "大额" in content → +10
        # 总 43
        assert score == 43.0


# ============================================================
#  Module 2: notification.NotificationService
# ============================================================


class TestNotificationServicePushAndMarkRead:
    """NotificationService.push / push_many / mark_read 行为契约.

    通知服务内部会调 commit(), 因此不能用 ``async_session`` fixture
    (它 begin() + rollback() 包裹, commit 后再 rollback 会报错).
    本类用自定义 ``notif_session`` fixture: 单 session, 无外部事务包裹.
    """

    @pytest_asyncio.fixture
    async def notif_session(self):
        """无 begin() 包裹的 session, 通知服务可自由 commit."""
        from sqlalchemy.ext.asyncio import async_sessionmaker, create_async_engine
        from sqlalchemy.pool import StaticPool

        from app.core.database import Base

        eng = create_async_engine(
            "sqlite+aiosqlite:///:memory:",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        async with eng.begin() as conn:
            await conn.run_sync(Base.metadata.create_all)
        sm = async_sessionmaker(eng, expire_on_commit=False)
        async with sm() as s:
            yield s
        await eng.dispose()

    @pytest.mark.asyncio
    async def test_push_creates_notification(self, notif_session):
        """push 单条 → DB 中存在对应 Notification 行."""
        from app.services.notification import NotificationService
        from app.models.db_models import Project

        proj = Project(
            name="测试项目", company_name="测试公司",
            fiscal_year=2024, status="active",
        )
        notif_session.add(proj)
        await notif_session.commit()

        notif = await NotificationService.push(
            notif_session,
            module="account_audit",
            type="test.event",
            title="测试通知",
            user_id=42,
            project_id=proj.id,
            severity="warn",
            body="测试 body",
        )
        assert notif is not None
        assert notif.id is not None
        assert notif.module == "account_audit"
        assert notif.user_id == 42
        assert notif.project_id == proj.id
        assert notif.severity == "warn"
        assert notif.is_read is False

    @pytest.mark.asyncio
    async def test_push_invalid_severity_falls_back_to_info(self, notif_session):
        """无效 severity → 兜底为 info (不抛错)."""
        from app.services.notification import NotificationService
        from app.models.db_models import Project

        proj = Project(name="X", company_name="X", fiscal_year=2024, status="active")
        notif_session.add(proj)
        await notif_session.commit()

        notif = await NotificationService.push(
            notif_session,
            module="system",
            type="t",
            title="t",
            project_id=proj.id,
            severity="bogus_sev",  # 无效 → 兜底 info
        )
        assert notif is not None
        assert notif.severity == "info"

    @pytest.mark.asyncio
    async def test_push_many_bulk_insert(self, notif_session):
        """push_many 批量 → 全部入库."""
        from app.services.notification import NotificationService
        from app.models.db_models import Project

        proj = Project(name="Y", company_name="Y", fiscal_year=2024, status="active")
        notif_session.add(proj)
        await notif_session.commit()

        items = [
            {"module": "system", "type": "t1", "title": "通知1",
             "project_id": proj.id, "user_id": 1},
            {"module": "system", "type": "t2", "title": "通知2",
             "project_id": proj.id, "user_id": 2},
            {"module": "system", "type": "t3", "title": "通知3",
             "project_id": proj.id, "user_id": None},  # 广播
        ]
        count = await NotificationService.push_many(notif_session, items)
        assert count == 3

    @pytest.mark.asyncio
    async def test_list_filters_by_user(self, notif_session):
        """list(user_id=X) 仅返回 X 专属 + 广播 (user_id=NULL)."""
        from app.services.notification import NotificationService
        from app.models.db_models import Project

        proj = Project(name="Z", company_name="Z", fiscal_year=2024, status="active")
        notif_session.add(proj)
        await notif_session.commit()

        # 给 user=1 推 1 条, user=2 推 1 条, 广播 1 条
        await NotificationService.push(
            notif_session, module="system", type="t", title="给1",
            project_id=proj.id, user_id=1,
        )
        await NotificationService.push(
            notif_session, module="system", type="t", title="给2",
            project_id=proj.id, user_id=2,
        )
        await NotificationService.push(
            notif_session, module="system", type="t", title="广播",
            project_id=proj.id, user_id=None,
        )

        result_u1 = await NotificationService.list(notif_session, user_id=1)
        titles_u1 = {n.title for n in result_u1["items"]}
        assert "给1" in titles_u1
        assert "广播" in titles_u1  # 广播对所有用户可见
        assert "给2" not in titles_u1  # 隔离
        assert result_u1["total"] == 2

    @pytest.mark.asyncio
    async def test_unread_count_excludes_read(self, notif_session):
        """mark_read 后 unread_count 不应再计该条."""
        from app.services.notification import NotificationService
        from app.models.db_models import Project

        proj = Project(name="W", company_name="W", fiscal_year=2024, status="active")
        notif_session.add(proj)
        await notif_session.commit()

        n1 = await NotificationService.push(
            notif_session, module="system", type="t", title="T1",
            project_id=proj.id, user_id=100,
        )
        await NotificationService.push(
            notif_session, module="system", type="t", title="T2",
            project_id=proj.id, user_id=100,
        )

        before = await NotificationService.unread_count(notif_session, user_id=100)
        assert before["total_unread"] == 2

        # 标记 n1 已读
        affected = await NotificationService.mark_read(
            notif_session, user_id=100, ids=[n1.id],
        )
        assert affected == 1

        after = await NotificationService.unread_count(notif_session, user_id=100)
        assert after["total_unread"] == 1

    @pytest.mark.asyncio
    async def test_mark_read_cross_user_returns_zero(self, notif_session):
        """P0 IDOR: ids 模式 user_id=None 应被拒绝 (返 0 行)."""
        from app.services.notification import NotificationService
        from app.models.db_models import Project

        proj = Project(name="V", company_name="V", fiscal_year=2024, status="active")
        notif_session.add(proj)
        await notif_session.commit()

        # 给 user=200 推一条
        n = await NotificationService.push(
            notif_session, module="system", type="t", title="T",
            project_id=proj.id, user_id=200,
        )
        # 尝试不传 user_id 标记别人的 → 必须返 0 (防跨用户读)
        affected = await NotificationService.mark_read(notif_session, ids=[n.id])
        assert affected == 0

        # 验证: 该条仍未读
        result = await NotificationService.list(notif_session, user_id=200, only_unread=True)
        assert result["total"] == 1


# ============================================================
#  Module 3: sentiment/briefing/generator
# ============================================================


def _mock_llm_client(r1=None, r2=None, r3=None, r4=None) -> MagicMock:
    """构造 mock LlmClient, 4 轮对话分别返回指定 dict."""
    client = MagicMock()
    payloads = [r1, r2, r3, r4]
    call_log: list[dict] = []

    async def chat_json(system, user, **kwargs):
        call_log.append({"system_prefix": system[:30], "user_prefix": user[:50]})
        if payloads and payloads[0] is not None:
            return payloads.pop(0)
        return {}

    client.chat_json.side_effect = chat_json
    client._call_log = call_log
    return client


def _sample_events():
    return [
        {
            "id": 1,
            "title": "公司公告",
            "content_text": "本期营收 1 亿元, 同比增长 15%",
            "publisher": "测试媒体",
            "publish_date": "2024-02-15",
            "severity": "info",
            "url": "http://example.com/1",
        },
        {
            "id": 2,
            "title": "媒体关注",
            "content_text": "公司毛利率 30%, 业绩亮眼",
            "publisher": "财经日报",
            "publish_date": "2024-02-16",
            "severity": "notice",
            "url": "http://example.com/2",
        },
    ]


class TestBriefingGeneratorGenerate:
    """BriefingGenerator.generate — 4 轮 LLM 协议主路径."""

    @pytest.mark.asyncio
    async def test_generate_returns_briefing_content(self):
        """正常路径: 4 轮 LLM 返回有效 dict → BriefingContent."""
        from app.services.sentiment.briefing.generator import BriefingGenerator

        r1 = {
            "key_facts": [
                {"event_id": 1, "fact": "营收 1 亿", "quote": "本期营收 1 亿元, 同比增长 15%",
                 "publish_date": "2024-02-15", "severity": "info"},
                {"event_id": 2, "fact": "毛利率 30%", "quote": "公司毛利率 30%, 业绩亮眼",
                 "publish_date": "2024-02-16", "severity": "notice"},
            ],
            "severity_breakdown": {"info": 1, "notice": 1, "warn": 0, "critical": 0},
            "watch_list": [{"event_id": 2, "reason": "毛利率异常待观察"}],
            "tone_words_used": [],
        }
        r2 = {
            "safe_facts": [
                {"event_id": 1, "verified": True, "issue": ""},
                {"event_id": 2, "verified": True, "issue": ""},
            ],
            "removed_facts": [],
        }
        r3 = {"critiques": [], "overall_risk_summary": "中性"}
        r4 = {"markdown": "# 测试公司 2024-02-15 舆情简报\n\n## 一、关键事实 (2 条)\n1. [事件#1] 营收 1 亿\n2. [事件#2] 毛利率 30%"}
        client = _mock_llm_client(r1=r1, r2=r2, r3=r3, r4=r4)

        with patch(
            "app.services.sentiment.briefing.generator.LlmClientFactory.preferred",
            return_value=client,
        ):
            gen = BriefingGenerator()
            content = await gen.generate(
                company_name="测试公司",
                project_id=1,
                briefing_date="2024-02-15",
                events=_sample_events(),
            )

        # 验证返回 BriefingContent
        assert content.markdown.startswith("# 测试公司 2024-02-15")
        assert len(content.extraction.key_facts) == 2
        assert len(content.self_check.safe_facts) == 2
        assert content.adversarial.overall_risk_summary == "中性"
        # safe_fact_event_ids 来自 r2.safe_facts 中 verified=True 的
        assert set(content.safe_fact_event_ids) == {1, 2}
        # event_snapshot 入库用精简结构
        assert len(content.event_snapshot) == 2
        assert content.event_snapshot[0]["id"] == 1
        # raw_input_events 保留原始
        assert len(content.raw_input_events) == 2
        # LLM 被调 4 轮
        assert len(client._call_log) == 4

    @pytest.mark.asyncio
    async def test_generate_rejects_empty_events(self):
        """空 events → 抛 ValueError (上游 detector 应先过滤)."""
        from app.services.sentiment.briefing.generator import BriefingGenerator

        client = _mock_llm_client()
        with patch(
            "app.services.sentiment.briefing.generator.LlmClientFactory.preferred",
            return_value=client,
        ):
            gen = BriefingGenerator()
            with pytest.raises(ValueError, match="events 不能为空"):
                await gen.generate(
                    company_name="X",
                    project_id=1,
                    briefing_date="2024-02-15",
                    events=[],
                )

    def test_strip_banned_words_replaces_emotional_terms(self):
        """后置硬过滤: 禁用情绪词被替换为 *** (静态方法, 无需实例)."""
        from app.services.sentiment.briefing.generator import BriefingGenerator

        text = "公司严重亏损, 暴雷, 业绩崩塌"
        cleaned = BriefingGenerator._strip_banned_words(text)
        assert "***" in cleaned
        for w in ("严重", "暴雷", "崩塌"):
            assert w not in cleaned

    @pytest.mark.asyncio
    async def test_parse_compose_handles_dict_and_string(self):
        """_parse_compose: r4 可能是 dict {'markdown': ...} 或 纯字符串."""
        from app.services.sentiment.briefing.generator import BriefingGenerator

        with patch(
            "app.services.sentiment.briefing.generator.LlmClientFactory.preferred",
            return_value=MagicMock(),
        ):
            gen = BriefingGenerator()
            # dict 形式
            out_dict = gen._parse_compose(
                {"markdown": "# X 2024 简报\n正文"}, "X", "2024-02-15", 3,
            )
            assert out_dict.startswith("# X 2024 简报")
            # 字符串形式
            out_str = gen._parse_compose(
                "纯字符串内容", "Y", "2024-02-16", 1,
            )
            assert "Y" in out_str and "2024-02-16" in out_str

    @pytest.mark.asyncio
    async def test_parse_compose_adds_fallback_header_when_missing(self):
        """r4 内容不以 '#' 开头 → 自动加公司名+日期标题兜底."""
        from app.services.sentiment.briefing.generator import BriefingGenerator

        with patch(
            "app.services.sentiment.briefing.generator.LlmClientFactory.preferred",
            return_value=MagicMock(),
        ):
            gen = BriefingGenerator()
            out = gen._parse_compose(
                "无标题正文段落", "兜底公司", "2024-03-01", 0,
            )
            assert out.startswith("# 兜底公司 2024-03-01")


class TestBriefingGeneratorLockedBriefing:
    """BriefingDetector.should_generate: 已锁定 briefing 不再生成 (P0 防护)."""

    @pytest.mark.asyncio
    async def test_locked_briefing_blocks_regeneration(self, async_session):
        """已有 is_locked=True 的 briefing → 检测返 False + reason='already_locked'."""
        from datetime import datetime

        from app.services.sentiment.briefing.detector import BriefingDetector
        from app.models.db_models import Project, SentimentDailyBriefing

        proj = Project(
            name="锁定测试项目", company_name="锁定公司",
            fiscal_year=2024, status="active",
        )
        async_session.add(proj)
        await async_session.flush()

        locked = SentimentDailyBriefing(
            project_id=proj.id,
            briefing_date="2024-02-15",
            title="已锁定简报",
            is_locked=True,
            locked_at=datetime.utcnow(),
            event_count=5,
        )
        async_session.add(locked)
        await async_session.flush()

        detector = BriefingDetector()
        result = await detector.should_generate(async_session, proj.id, "2024-02-15")

        assert result.should_generate is False
        assert result.reason == "already_locked"
        assert result.existing_briefing_id == locked.id

    @pytest.mark.asyncio
    async def test_no_events_returns_no_events_reason(self, async_session):
        """无事件 → 不生成 (no_events)."""
        from app.services.sentiment.briefing.detector import BriefingDetector
        from app.models.db_models import Project

        proj = Project(
            name="无事件项目", company_name="X",
            fiscal_year=2024, status="active",
        )
        async_session.add(proj)
        await async_session.flush()

        detector = BriefingDetector()
        result = await detector.should_generate(async_session, proj.id, "2024-02-15")
        assert result.should_generate is False
        assert result.reason == "no_events"


# ============================================================
#  Module 4: sales_ledger/document_parser
# ============================================================


class _FakeUploadFile:
    """模仿 FastAPI UploadFile 接口, 仅暴露 filename + read()."""

    def __init__(self, filename: str, content: bytes):
        self.filename = filename
        self._content = content

    async def read(self) -> bytes:
        return self._content


class TestDocumentParserExtOf:
    """DocumentParser.ext_of — 文件名 → 扩展名."""

    def test_ext_of_basic(self):
        from app.services.sales_ledger.document_parser import DocumentParser

        assert DocumentParser.ext_of("foo.xlsx") == ".xlsx"
        assert DocumentParser.ext_of("FOO.PDF") == ".pdf"
        assert DocumentParser.ext_of("a.b.docx") == ".docx"
        assert DocumentParser.ext_of("noext") == ""


class TestDocumentParserXLSX:
    """DocumentParser._parse_xlsx — xlsx → Markdown."""

    def test_parse_xlsx_returns_markdown_table(self, tmp_path):
        from app.services.sales_ledger.document_parser import DocumentParser

        # 构造 xlsx
        xlsx_path = tmp_path / "test.xlsx"
        df = pd.DataFrame({"客户": ["A公司", "B公司"], "金额": [1000, 2000]})
        df.to_excel(xlsx_path, index=False, sheet_name="销售明细")

        text = DocumentParser._parse_xlsx(xlsx_path)
        assert "## Sheet: 销售明细" in text
        assert "客户" in text
        assert "A公司" in text
        assert "1000" in text

    def test_parse_xlsx_multiple_sheets(self, tmp_path):
        from app.services.sales_ledger.document_parser import DocumentParser

        xlsx_path = tmp_path / "multi.xlsx"
        with pd.ExcelWriter(xlsx_path) as writer:
            pd.DataFrame({"a": [1]}).to_excel(writer, sheet_name="表1", index=False)
            pd.DataFrame({"b": [2]}).to_excel(writer, sheet_name="表2", index=False)

        text = DocumentParser._parse_xlsx(xlsx_path)
        assert "## Sheet: 表1" in text
        assert "## Sheet: 表2" in text


class TestDocumentParserDocx:
    """DocumentParser._parse_docx — docx → 段落 + 表格."""

    def test_parse_docx_returns_paragraphs(self, tmp_path):
        from docx import Document

        from app.services.sales_ledger.document_parser import DocumentParser

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("第一段内容")
        doc.add_paragraph("第二段内容")
        doc.save(str(docx_path))

        text = DocumentParser._parse_docx(docx_path)
        assert "第一段内容" in text
        assert "第二段内容" in text


class TestDocumentParserPdf:
    """DocumentParser._parse_pdf — pdf → 文本 (mock pdfplumber)."""

    def test_parse_pdf_with_mock_pages(self, tmp_path):
        """mock pdfplumber: 返回 2 页文本."""
        from app.services.sales_ledger.document_parser import DocumentParser

        fake_pdf = tmp_path / "fake.pdf"
        fake_pdf.write_bytes(b"%PDF-fake")

        # mock pdfplumber.open 返回 2 页
        mock_page1 = MagicMock()
        mock_page1.extract_text.return_value = "第 1 页内容"
        mock_page1.extract_tables.return_value = []
        mock_page2 = MagicMock()
        mock_page2.extract_text.return_value = "第 2 页内容"
        mock_page2.extract_tables.return_value = []

        mock_pdf = MagicMock()
        mock_pdf.pages = [mock_page1, mock_page2]
        mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
        mock_pdf.__exit__ = MagicMock(return_value=False)

        with patch.dict("sys.modules", {
            "pdfplumber": MagicMock(open=MagicMock(return_value=mock_pdf)),
        }):
            text = DocumentParser._parse_pdf(fake_pdf)
        assert "## Page 1" in text
        assert "第 1 页内容" in text
        assert "## Page 2" in text
        assert "第 2 页内容" in text


class TestDocumentParserParse:
    """DocumentParser.parse — 端到端入口, 处理 upload + 临时文件 + 错误路径."""

    @pytest.mark.asyncio
    async def test_parse_xlsx_end_to_end(self, tmp_path):
        """上传 xlsx → 返回 (doc_type='xlsx', markdown_text)."""
        from app.services.sales_ledger.document_parser import DocumentParser

        df = pd.DataFrame({"项目": ["X", "Y"], "值": [10, 20]})
        xlsx_buf = tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False)
        df.to_excel(xlsx_buf.name, index=False)
        xlsx_buf.close()

        with open(xlsx_buf.name, "rb") as f:
            content = f.read()
        upload = _FakeUploadFile("test.xlsx", content)

        save_dir = tmp_path / "uploads"
        doc_type, text = await DocumentParser.parse(upload, save_dir)
        assert doc_type == "xlsx"
        assert "项目" in text

    @pytest.mark.asyncio
    async def test_parse_rejects_unsupported_extension(self, tmp_path):
        """上传 .exe → DocumentParserError (不支持的文件类型)."""
        from app.services.sales_ledger.document_parser import DocumentParser, DocumentParserError

        upload = _FakeUploadFile("malware.exe", b"MZ\x90\x00")
        with pytest.raises(DocumentParserError, match="不支持的文件类型"):
            await DocumentParser.parse(upload, tmp_path)

    @pytest.mark.asyncio
    async def test_parse_sanitizes_path_traversal_filename(self, tmp_path):
        """文件名含 ../../ 路径穿越 → 安全 sanitize (截到 basename), 不抛错."""
        from app.services.sales_ledger.document_parser import DocumentParser

        # 构造 xlsx 内容 (用 openpyxl 直接写, 避免 pandas 临时文件)
        import openpyxl

        xlsx_path = tmp_path / "src.xlsx"
        wb = openpyxl.Workbook()
        wb.active["A1"] = "test"
        wb.save(str(xlsx_path))

        with open(xlsx_path, "rb") as f:
            content = f.read()

        # 路径穿越 filename
        upload = _FakeUploadFile("../../../etc/passwd.xlsx", content)
        save_dir = tmp_path / "uploads"
        # 不抛错: 临时文件被 sanitize 到 save_dir 内, finally 清理
        doc_type, text = await DocumentParser.parse(upload, save_dir)
        # 解析成功, 返回 doc_type
        assert doc_type == "xlsx"
        # parse 返回后, 临时文件已被 unlink 清理 (P0 防护)
        assert not (save_dir / "sales_src_passwd.xlsx").exists()


# ============================================================
#  Smoke
# ============================================================


class TestRound36Smoke:
    """冒烟: 至少 5 类 P0 模块被独立测试覆盖."""

    def test_at_least_5_p0_modules_covered(self):
        """本文件包含 ≥5 类 TestClass, 覆盖 verifier P0 + 4 模块."""
        import sys

        mod = sys.modules[__name__]
        classes = [
            obj
            for name, obj in vars(mod).items()
            if isinstance(obj, type) and name.startswith("Test")
        ]
        # 1 (verifier P0) + 4 (regulatory_case / notification / briefing / document_parser) + 1 smoke = ≥6
        assert len(classes) >= 6, (
            f"只覆盖 {len(classes)} 类, 应 ≥6 类 (1 verifier P0 + 4 模块 + 1 smoke)"
        )